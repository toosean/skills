#!/usr/bin/env python3
"""Convert Markdown files to Word documents with a pure Python pipeline."""

from __future__ import annotations

import argparse
import io
import json
import sys
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from urllib.error import HTTPError, URLError
from urllib.parse import unquote, urlparse
from urllib.request import Request, urlopen

try:
    from markdown_it import MarkdownIt
except ImportError as exc:  # pragma: no cover - exercised by CLI dependency checks
    raise SystemExit(
        "Missing dependency: markdown-it-py. Install it with `python -m pip install markdown-it-py`."
    ) from exc

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.shared import Cm, Inches, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - exercised by CLI dependency checks
    raise SystemExit(
        "Missing dependency: python-docx. Install it with `python -m pip install python-docx`."
    ) from exc


BODY_EAST_ASIA_FONT = "SimSun"
BODY_LATIN_FONT = "Times New Roman"
HEADING_EAST_ASIA_FONT = "SimHei"
CODE_FONT = "Consolas"
MAX_IMAGE_WIDTH = Inches(5.8)
MAX_REMOTE_IMAGE_BYTES = 20 * 1024 * 1024
REMOTE_IMAGE_TIMEOUT_SECONDS = 20
REMOTE_IMAGE_USER_AGENT = "markdown-to-word/1.0"


@dataclass
class RenderContext:
    list_level: int = 0
    blockquote_depth: int = 0
    continuation_indent: int = 0


@dataclass
class InlineState:
    bold: bool = False
    italic: bool = False
    code: bool = False
    href: str | None = None


def clone_inline_state(state: InlineState) -> InlineState:
    return InlineState(
        bold=state.bold,
        italic=state.italic,
        code=state.code,
        href=state.href,
    )


def attr_get(token, name: str, default: str | None = None) -> str | None:
    value = token.attrGet(name)
    return value if value is not None else default


def clean_cli_path(value: str | Path) -> Path:
    text = str(value)
    cleaned = "".join(char for char in text if unicodedata.category(char) != "Cf").strip()
    return Path(cleaned)


def is_windows_absolute_path_text(value: str) -> bool:
    return (
        len(value) >= 3
        and value[0].isalpha()
        and value[1] == ":"
        and value[2] in {"\\", "/"}
    ) or value.startswith("\\\\")


def normalize_local_path_text(value: str) -> str:
    text = value.strip()
    if len(text) >= 4 and text[0] == "/" and text[1].isalpha() and text[2] == ":":
        return text[1:]
    return text


def parse_front_matter(markdown: str) -> tuple[dict[str, str], str]:
    text = markdown.lstrip("\ufeff")
    lines = text.splitlines()
    if not lines or lines[0].strip() != "---":
        return {}, text

    end_index = None
    for index in range(1, len(lines)):
        if lines[index].strip() in {"---", "..."}:
            end_index = index
            break

    if end_index is None:
        return {}, text

    metadata: dict[str, str] = {}
    for line in lines[1:end_index]:
        if not line.strip() or line.lstrip().startswith("#"):
            continue
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        key = key.strip().lower()
        value = value.strip().strip('"').strip("'")
        if key:
            metadata[key] = value

    body = "\n".join(lines[end_index + 1 :])
    if text.endswith("\n"):
        body += "\n"
    return metadata, body


def style_exists(document: Document, style_name: str) -> bool:
    try:
        document.styles[style_name]
        return True
    except KeyError:
        return False


def get_or_add_paragraph_style(document: Document, name: str, base: str = "Normal"):
    styles = document.styles
    try:
        return styles[name]
    except KeyError:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        if base in styles:
            style.base_style = styles[base]
        return style


def get_or_add_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def remove_children(parent, tag: str) -> None:
    for child in list(parent.findall(qn(tag))):
        parent.remove(child)


def set_run_east_asia_font(run_or_font, east_asia: str, ascii_font: str | None = None) -> None:
    element = run_or_font._element
    rpr = element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if ascii_font:
        rfonts.set(qn("w:ascii"), ascii_font)
        rfonts.set(qn("w:hAnsi"), ascii_font)


def set_style_font(
    style,
    east_asia: str,
    ascii_font: str | None = None,
    size: Pt | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    if ascii_font:
        style.font.name = ascii_font
    set_run_east_asia_font(style.font, east_asia, ascii_font)
    if size is not None:
        style.font.size = size
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    if color is not None:
        style.font.color.rgb = color


def set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    existing = ppr.find(qn("w:shd"))
    if existing is not None:
        ppr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    ppr.append(shading)


def set_run_shading(run, fill: str) -> None:
    rpr = run._r.get_or_add_rPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    rpr.append(shading)


def set_paragraph_bottom_border(paragraph, color: str = "808080", size: str = "6") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = pbdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pbdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    existing = tcpr.find(qn("w:shd"))
    if existing is not None:
        tcpr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tcpr.append(shading)


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    margins = tcpr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcpr.append(margins)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def add_update_fields_setting(document: Document) -> None:
    settings = document.settings.element
    if settings.find(qn("w:updateFields")) is not None:
        return
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def set_paragraph_style_id(paragraph, style_id: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pstyle = ppr.find(qn("w:pStyle"))
    if pstyle is None:
        pstyle = OxmlElement("w:pStyle")
        ppr.insert(0, pstyle)
    pstyle.set(qn("w:val"), style_id)


def add_toc_field(document: Document, depth: int) -> None:
    title = document.add_paragraph("目录", style="TOC Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = document.add_paragraph()
    set_paragraph_style_id(paragraph, "TOC1")
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\o "1-{depth}" \\h \\z \\u'

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "请在 Word 中更新域以生成目录"
    placeholder.append(placeholder_text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    paragraph._p.append(placeholder)
    final_run = paragraph.add_run()
    final_run._r.append(end)
    document.add_paragraph()
    add_update_fields_setting(document)


def apply_formal_zh_defaults(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    set_style_font(normal, BODY_EAST_ASIA_FONT, BODY_LATIN_FONT, Pt(12))
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)

    title = document.styles["Title"]
    set_style_font(title, HEADING_EAST_ASIA_FONT, BODY_LATIN_FONT, Pt(18), bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)

    for level in range(1, 7):
        style_name = f"Heading {level}"
        if style_name not in document.styles:
            continue
        style = document.styles[style_name]
        size = {1: 16, 2: 14, 3: 12.5}.get(level, 12)
        set_style_font(style, HEADING_EAST_ASIA_FONT, BODY_LATIN_FONT, Pt(size), bold=True)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    toc_title = get_or_add_paragraph_style(document, "TOC Title")
    set_style_font(toc_title, HEADING_EAST_ASIA_FONT, BODY_LATIN_FONT, Pt(16), bold=True)
    toc_title.paragraph_format.first_line_indent = Pt(0)
    toc_title.paragraph_format.space_after = Pt(10)
    configure_toc_styles(document)

    quote = get_or_add_paragraph_style(document, "Markdown Quote")
    set_style_font(quote, BODY_EAST_ASIA_FONT, BODY_LATIN_FONT, Pt(11), italic=False, color=RGBColor(80, 80, 80))
    quote.paragraph_format.left_indent = Cm(0.75)
    quote.paragraph_format.first_line_indent = Pt(0)
    quote.paragraph_format.space_before = Pt(4)
    quote.paragraph_format.space_after = Pt(4)

    code = get_or_add_paragraph_style(document, "Markdown Code Block")
    set_style_font(code, BODY_EAST_ASIA_FONT, CODE_FONT, Pt(9))
    code.paragraph_format.first_line_indent = Pt(0)
    code.paragraph_format.left_indent = Cm(0.4)
    code.paragraph_format.right_indent = Cm(0.2)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.0

    for style_name in (
        "List Bullet",
        "List Bullet 2",
        "List Bullet 3",
        "List Number",
        "List Number 2",
        "List Number 3",
    ):
        if style_exists(document, style_name):
            set_style_font(document.styles[style_name], BODY_EAST_ASIA_FONT, BODY_LATIN_FONT, Pt(12))


def configure_toc_styles(document: Document) -> None:
    section = document.sections[0]
    content_width = section.page_width - section.left_margin - section.right_margin

    for level in range(1, 7):
        style = get_or_add_paragraph_style(document, f"TOC {level}")
        style_element = style._element
        style_element.set(qn("w:styleId"), f"TOC{level}")
        style_element.attrib.pop(qn("w:customStyle"), None)
        remove_children(style_element, "w:basedOn")
        name = get_or_add_child(style_element, "w:name")
        name.set(qn("w:val"), f"toc {level}")
        ui_priority = get_or_add_child(style_element, "w:uiPriority")
        ui_priority.set(qn("w:val"), str(39 + level))
        get_or_add_child(style_element, "w:unhideWhenUsed")

        set_style_font(style, BODY_EAST_ASIA_FONT, BODY_LATIN_FONT, Pt(11))
        paragraph_format = style.paragraph_format
        paragraph_format.first_line_indent = Pt(0)
        paragraph_format.left_indent = Cm((level - 1) * 0.74)
        paragraph_format.right_indent = Pt(0)
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(4)
        paragraph_format.line_spacing = 1.15
        paragraph_format.tab_stops.clear_all()
        paragraph_format.tab_stops.add_tab_stop(
            content_width,
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )

        ppr = style_element.get_or_add_pPr()
        ind = get_or_add_child(ppr, "w:ind")
        ind.set(qn("w:left"), str(int(Cm((level - 1) * 0.74).twips)))
        ind.set(qn("w:firstLine"), "0")
        if qn("w:hanging") in ind.attrib:
            del ind.attrib[qn("w:hanging")]


def list_style_name(document: Document, ordered: bool, level: int) -> str:
    base = "List Number" if ordered else "List Bullet"
    if level <= 0:
        candidate = base
    else:
        candidate = f"{base} {min(level + 1, 3)}"
    if style_exists(document, candidate):
        return candidate
    return base


def add_hyperlink_run(paragraph, href: str, text: str, state: InlineState) -> None:
    if not text:
        return
    relationship_id = paragraph.part.relate_to(href, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), BODY_EAST_ASIA_FONT)
    rfonts.set(qn("w:ascii"), BODY_LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_LATIN_FONT)
    rpr.append(rfonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    if state.bold:
        rpr.append(OxmlElement("w:b"))
    if state.italic:
        rpr.append(OxmlElement("w:i"))

    run.append(rpr)
    text_node = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


class MarkdownDocxRenderer:
    def __init__(
        self,
        document: Document,
        input_path: Path,
        resource_paths: Iterable[Path] | None = None,
        number_headings: bool = True,
        heading_level_offset: int = 0,
    ) -> None:
        self.document = document
        self.input_path = input_path.resolve()
        self.resource_paths = [self.input_path.parent]
        if resource_paths:
            for resource_path in resource_paths:
                if resource_path.is_absolute():
                    self.resource_paths.append(resource_path)
                else:
                    self.resource_paths.append((self.input_path.parent / resource_path).resolve())
        self.number_headings = number_headings
        self.heading_level_offset = heading_level_offset
        self.heading_counters = [0, 0, 0, 0, 0, 0]
        self.warnings: list[str] = []

    def render(self, tokens) -> None:
        self.render_blocks(tokens, 0, len(tokens), RenderContext())

    def render_blocks(self, tokens, start: int, end: int, context: RenderContext) -> int:
        index = start
        while index < end:
            token = tokens[index]
            token_type = token.type

            if token_type == "heading_open":
                index = self.render_heading(tokens, index)
            elif token_type == "paragraph_open":
                index = self.render_paragraph(tokens, index, context)
            elif token_type in {"bullet_list_open", "ordered_list_open"}:
                index = self.render_list(tokens, index, context)
            elif token_type == "blockquote_open":
                close_index = find_matching_token(tokens, index, "blockquote_open", "blockquote_close")
                child_context = RenderContext(
                    list_level=context.list_level,
                    blockquote_depth=context.blockquote_depth + 1,
                    continuation_indent=context.continuation_indent,
                )
                self.render_blocks(tokens, index + 1, close_index, child_context)
                index = close_index + 1
            elif token_type in {"fence", "code_block"}:
                self.add_code_block(token.content.rstrip("\n"))
                index += 1
            elif token_type == "hr":
                self.add_horizontal_rule()
                index += 1
            elif token_type == "table_open":
                close_index = find_matching_token(tokens, index, "table_open", "table_close")
                self.render_table(tokens[index + 1 : close_index])
                index = close_index + 1
            elif token_type.endswith("_close"):
                return index
            else:
                index += 1
        return index

    def render_heading(self, tokens, index: int) -> int:
        token = tokens[index]
        markdown_level = int(token.tag[1]) if token.tag and token.tag.startswith("h") else 1
        level = max(1, min(6, markdown_level - self.heading_level_offset))
        inline = tokens[index + 1] if index + 1 < len(tokens) and tokens[index + 1].type == "inline" else None
        text_prefix = ""
        if self.number_headings:
            self.heading_counters[level - 1] += 1
            for offset in range(level, len(self.heading_counters)):
                self.heading_counters[offset] = 0
            number = ".".join(str(value) for value in self.heading_counters[:level] if value)
            text_prefix = f"{number} "

        paragraph = self.document.add_paragraph(style=f"Heading {min(level, 6)}")
        paragraph.paragraph_format.first_line_indent = Pt(0)
        if text_prefix:
            run = paragraph.add_run(text_prefix)
            run.bold = True
            self.apply_run_style(run, InlineState(bold=True))
        if inline is not None:
            self.render_inline(paragraph, inline.children or [], InlineState())
        return index + 3

    def render_paragraph(self, tokens, index: int, context: RenderContext) -> int:
        inline = tokens[index + 1] if index + 1 < len(tokens) and tokens[index + 1].type == "inline" else None
        paragraph = self.add_context_paragraph(context)
        if inline is not None:
            self.render_inline(paragraph, inline.children or [], InlineState())
        return index + 3

    def add_context_paragraph(self, context: RenderContext, style: str | None = None):
        if style is None:
            style = "Markdown Quote" if context.blockquote_depth else "Normal"
        paragraph = self.document.add_paragraph(style=style)
        if style == "Normal" and context.continuation_indent:
            paragraph.paragraph_format.left_indent = Cm(0.55 * context.continuation_indent)
            paragraph.paragraph_format.first_line_indent = Pt(0)
        if context.blockquote_depth and style != "Markdown Quote":
            paragraph.paragraph_format.left_indent = Cm(0.75 * context.blockquote_depth)
        return paragraph

    def render_list(self, tokens, index: int, context: RenderContext) -> int:
        list_token = tokens[index]
        ordered = list_token.type == "ordered_list_open"
        close_type = "ordered_list_close" if ordered else "bullet_list_close"
        close_index = find_matching_token(tokens, index, list_token.type, close_type)
        child_index = index + 1
        while child_index < close_index:
            if tokens[child_index].type == "list_item_open":
                item_close = find_matching_token(tokens, child_index, "list_item_open", "list_item_close")
                self.render_list_item(tokens, child_index + 1, item_close, ordered, context)
                child_index = item_close + 1
            else:
                child_index += 1
        return close_index + 1

    def render_list_item(self, tokens, start: int, end: int, ordered: bool, context: RenderContext) -> None:
        index = start
        level = context.list_level
        first_block_rendered = False

        while index < end:
            token = tokens[index]
            if token.type == "paragraph_open" and not first_block_rendered:
                inline = tokens[index + 1] if index + 1 < end and tokens[index + 1].type == "inline" else None
                style_name = list_style_name(self.document, ordered, level)
                paragraph = self.document.add_paragraph(style=style_name)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                if context.blockquote_depth:
                    paragraph.paragraph_format.left_indent = Cm(0.75 * context.blockquote_depth + 0.45 * level)
                if inline is not None:
                    self.render_inline(paragraph, inline.children or [], InlineState())
                first_block_rendered = True
                index += 3
            elif token.type in {"bullet_list_open", "ordered_list_open"}:
                nested_context = RenderContext(
                    list_level=level + 1,
                    blockquote_depth=context.blockquote_depth,
                    continuation_indent=level + 1,
                )
                index = self.render_list(tokens, index, nested_context)
                first_block_rendered = True
            else:
                nested_context = RenderContext(
                    list_level=level + 1,
                    blockquote_depth=context.blockquote_depth,
                    continuation_indent=level + 1,
                )
                next_index = self.render_blocks(tokens, index, end, nested_context)
                if next_index == index:
                    index += 1
                else:
                    index = next_index
                first_block_rendered = True

        if not first_block_rendered:
            style_name = list_style_name(self.document, ordered, level)
            self.document.add_paragraph("", style=style_name)

    def render_inline(self, paragraph, children, state: InlineState) -> None:
        for child in children:
            token_type = child.type
            if token_type in {"text", "text_special"}:
                self.add_text(paragraph, child.content, state)
            elif token_type == "code_inline":
                code_state = clone_inline_state(state)
                code_state.code = True
                self.add_text(paragraph, child.content, code_state)
            elif token_type == "softbreak":
                paragraph.add_run().add_break()
            elif token_type == "hardbreak":
                paragraph.add_run().add_break(WD_BREAK.LINE)
            elif token_type == "strong_open":
                state = clone_inline_state(state)
                state.bold = True
            elif token_type == "strong_close":
                state = clone_inline_state(state)
                state.bold = False
            elif token_type == "em_open":
                state = clone_inline_state(state)
                state.italic = True
            elif token_type == "em_close":
                state = clone_inline_state(state)
                state.italic = False
            elif token_type == "link_open":
                state = clone_inline_state(state)
                state.href = attr_get(child, "href")
            elif token_type == "link_close":
                state = clone_inline_state(state)
                state.href = None
            elif token_type == "image":
                self.add_image(paragraph, child)
            elif token_type == "html_inline":
                self.add_text(paragraph, child.content, state)

    def add_text(self, paragraph, text: str, state: InlineState) -> None:
        if not text:
            return
        if state.href and not state.code:
            add_hyperlink_run(paragraph, state.href, text, state)
            return
        run = paragraph.add_run(text)
        self.apply_run_style(run, state)

    def apply_run_style(self, run, state: InlineState) -> None:
        if state.bold:
            run.bold = True
        if state.italic:
            run.italic = True
        if state.code:
            run.font.name = CODE_FONT
            run.font.size = Pt(9)
            set_run_east_asia_font(run, BODY_EAST_ASIA_FONT, CODE_FONT)
            set_run_shading(run, "F2F2F2")
        else:
            set_run_east_asia_font(run, BODY_EAST_ASIA_FONT, BODY_LATIN_FONT)

    def add_image(self, paragraph, token) -> None:
        src = attr_get(token, "src")
        alt = token.content or src or "image"
        if not src:
            self.add_text(paragraph, f"[image: {alt}]", InlineState())
            return

        image_source = self.resolve_image_source(src)
        if image_source is None:
            self.warnings.append(f"Image not found: {src}")
            self.add_text(paragraph, f"[image not found: {alt}]", InlineState(italic=True))
            return

        try:
            run = paragraph.add_run()
            if isinstance(image_source, Path):
                run.add_picture(str(image_source), width=MAX_IMAGE_WIDTH)
            else:
                image_source.seek(0)
                run.add_picture(image_source, width=MAX_IMAGE_WIDTH)
        except Exception as exc:  # pragma: no cover - depends on image codec support
            self.warnings.append(f"Could not insert image {src}: {exc}")
            self.add_text(paragraph, f"[image could not be inserted: {alt}]", InlineState(italic=True))

    def resolve_image_source(self, src: str) -> Path | io.BytesIO | None:
        src = src.strip()
        decoded_src = normalize_local_path_text(unquote(src))
        if is_windows_absolute_path_text(src) or is_windows_absolute_path_text(decoded_src):
            return self.resolve_local_image_path(decoded_src)

        parsed = urlparse(src)
        if parsed.scheme in {"http", "https"}:
            return self.download_remote_image(src)
        if parsed.scheme == "file":
            file_path = normalize_local_path_text(unquote(parsed.path))
            if parsed.netloc:
                file_path = f"//{parsed.netloc}{file_path}"
            return self.resolve_local_image_path(file_path)
        if parsed.scheme:
            self.warnings.append(f"Unsupported image URI scheme: {parsed.scheme}")
            return None
        local_candidates = [decoded_src]
        if parsed.path:
            local_candidates.append(normalize_local_path_text(unquote(parsed.path)))
        for local_src in dict.fromkeys(local_candidates):
            resolved = self.resolve_local_image_path(local_src)
            if resolved is not None:
                return resolved
        return None

    def resolve_local_image_path(self, src: str) -> Path | None:
        candidate = Path(src)
        if candidate.is_absolute() and candidate.exists():
            return candidate
        for base in self.resource_paths:
            resolved = (base / src).resolve()
            if resolved.exists():
                return resolved
        return None

    def download_remote_image(self, src: str) -> io.BytesIO | None:
        request = Request(src, headers={"User-Agent": REMOTE_IMAGE_USER_AGENT})
        try:
            with urlopen(request, timeout=REMOTE_IMAGE_TIMEOUT_SECONDS) as response:
                content_length = response.headers.get("Content-Length")
                if content_length and int(content_length) > MAX_REMOTE_IMAGE_BYTES:
                    self.warnings.append(
                        f"Remote image too large ({content_length} bytes, max {MAX_REMOTE_IMAGE_BYTES}): {src}"
                    )
                    return None

                data = bytearray()
                while True:
                    chunk = response.read(1024 * 64)
                    if not chunk:
                        break
                    data.extend(chunk)
                    if len(data) > MAX_REMOTE_IMAGE_BYTES:
                        self.warnings.append(
                            f"Remote image too large (over {MAX_REMOTE_IMAGE_BYTES} bytes): {src}"
                        )
                        return None

                content_type = response.headers.get("Content-Type", "")
                if content_type and not content_type.lower().startswith("image/"):
                    self.warnings.append(f"Remote URL did not return an image content type ({content_type}): {src}")

                return io.BytesIO(bytes(data))
        except (HTTPError, URLError, TimeoutError, ValueError, OSError) as exc:
            self.warnings.append(f"Could not download remote image {src}: {exc}")
            return None

    def add_code_block(self, content: str) -> None:
        lines = content.splitlines() or [""]
        for line in lines:
            paragraph = self.document.add_paragraph(style="Markdown Code Block")
            paragraph.paragraph_format.first_line_indent = Pt(0)
            set_paragraph_shading(paragraph, "F7F7F7")
            run = paragraph.add_run(line if line else " ")
            run.font.name = CODE_FONT
            run.font.size = Pt(9)
            set_run_east_asia_font(run, BODY_EAST_ASIA_FONT, CODE_FONT)

    def add_horizontal_rule(self) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(0)
        set_paragraph_bottom_border(paragraph)

    def render_table(self, tokens) -> None:
        rows: list[list[list]] = []
        header_rows: set[int] = set()
        current_row: list[list] | None = None
        current_cell: list | None = None
        current_cell_is_header = False

        for token in tokens:
            if token.type == "tr_open":
                current_row = []
            elif token.type == "tr_close":
                if current_row is not None:
                    rows.append(current_row)
                current_row = None
            elif token.type in {"th_open", "td_open"}:
                current_cell = []
                current_cell_is_header = token.type == "th_open"
            elif token.type in {"th_close", "td_close"}:
                if current_row is not None and current_cell is not None:
                    current_row.append(current_cell)
                    if current_cell_is_header:
                        header_rows.add(len(rows))
                current_cell = None
                current_cell_is_header = False
            elif token.type == "inline" and current_cell is not None:
                current_cell.extend(token.children or [])

        if not rows:
            return

        column_count = max(len(row) for row in rows)
        table = self.document.add_table(rows=len(rows), cols=column_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        if style_exists(self.document, "Table Grid"):
            table.style = "Table Grid"

        for row_index, row_cells in enumerate(rows):
            row = table.rows[row_index]
            if row_index in header_rows:
                set_repeat_table_header(row)
            for column_index in range(column_count):
                cell = row.cells[column_index]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                if column_index < len(row_cells):
                    state = InlineState(bold=row_index in header_rows)
                    self.render_inline(paragraph, row_cells[column_index], state)
                if row_index in header_rows:
                    set_cell_shading(cell, "EDEDED")
                    for run in paragraph.runs:
                        run.bold = True


def find_matching_token(tokens, start: int, open_type: str, close_type: str) -> int:
    depth = 1
    for index in range(start + 1, len(tokens)):
        token_type = tokens[index].type
        if token_type == open_type:
            depth += 1
        elif token_type == close_type:
            depth -= 1
            if depth == 0:
                return index
    raise ValueError(f"Could not find matching {close_type} for {open_type}")


def build_markdown_parser() -> MarkdownIt:
    parser = MarkdownIt("commonmark", {"html": False})
    parser.enable("table")
    parser.enable("strikethrough")
    return parser


def heading_level_offset(tokens) -> int:
    levels = [
        int(token.tag[1])
        for token in tokens
        if token.type == "heading_open" and token.tag and token.tag.startswith("h") and token.tag[1:].isdigit()
    ]
    if not levels:
        return 0
    return max(0, min(levels) - 1)


def add_metadata(document: Document, metadata: dict[str, str]) -> None:
    title = metadata.get("title")
    if title:
        document.core_properties.title = title
        document.add_paragraph(title, style="Title")

    details = []
    if metadata.get("author"):
        document.core_properties.author = metadata["author"]
        details.append(metadata["author"])
    if metadata.get("date"):
        details.append(metadata["date"])

    for value in details:
        paragraph = document.add_paragraph(value)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        for run in paragraph.runs:
            set_run_east_asia_font(run, BODY_EAST_ASIA_FONT, BODY_LATIN_FONT)

    if title or details:
        document.add_paragraph()


def convert_markdown(
    input_path: Path,
    output_path: Path,
    *,
    format_name: str = "formal-zh",
    toc: bool = True,
    number_headings: bool = True,
    toc_depth: int = 3,
    resource_paths: Iterable[Path] | None = None,
) -> dict[str, object]:
    if format_name != "formal-zh":
        raise ValueError(f"Unsupported format: {format_name}")
    if not input_path.exists():
        raise FileNotFoundError(f"Input file does not exist: {input_path}")
    if input_path.suffix.lower() != ".md":
        raise ValueError("Input file must have a .md extension")
    if toc_depth < 1 or toc_depth > 6:
        raise ValueError("--toc-depth must be between 1 and 6")

    markdown = input_path.read_text(encoding="utf-8")
    metadata, body = parse_front_matter(markdown)

    document = Document()
    apply_formal_zh_defaults(document)
    add_metadata(document, metadata)
    if toc:
        add_toc_field(document, toc_depth)

    parser = build_markdown_parser()
    tokens = parser.parse(body)
    heading_offset = heading_level_offset(tokens)
    renderer = MarkdownDocxRenderer(
        document,
        input_path=input_path,
        resource_paths=resource_paths,
        number_headings=number_headings,
        heading_level_offset=heading_offset,
    )
    renderer.render(tokens)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)

    return {
        "ok": True,
        "input": str(input_path),
        "output": str(output_path),
        "format": format_name,
        "toc": toc,
        "number_headings": number_headings,
        "toc_depth": toc_depth,
        "heading_level_offset": heading_offset,
        "warnings": renderer.warnings,
    }


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Convert Markdown files to Word documents.")
    subparsers = parser.add_subparsers(dest="command", required=True)

    convert = subparsers.add_parser("convert", help="Convert a Markdown file to .docx.")
    convert.add_argument("input", help="Path to the input .md file.")
    convert.add_argument("--output", "-o", required=True, help="Path to the output .docx file.")
    convert.add_argument("--format", default="formal-zh", choices=["formal-zh"], help="Built-in output format.")
    toc_group = convert.add_mutually_exclusive_group()
    toc_group.add_argument("--toc", dest="toc", action="store_true", default=True, help="Insert a Word TOC field.")
    toc_group.add_argument("--no-toc", dest="toc", action="store_false", help="Do not insert a TOC field.")
    numbering_group = convert.add_mutually_exclusive_group()
    numbering_group.add_argument(
        "--number-headings",
        dest="number_headings",
        action="store_true",
        default=True,
        help="Prefix headings with visible section numbers.",
    )
    numbering_group.add_argument(
        "--no-number-headings",
        dest="number_headings",
        action="store_false",
        help="Keep heading text unnumbered.",
    )
    convert.add_argument("--toc-depth", type=int, default=3, help="Maximum heading depth included in the TOC field.")
    convert.add_argument(
        "--resource-path",
        action="append",
        default=[],
        help="Additional directory to search for relative image paths. Repeatable.",
    )
    convert.add_argument("--json", action="store_true", help="Print machine-readable conversion result.")
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_arg_parser()
    args = parser.parse_args(argv)

    try:
        if args.command == "convert":
            input_path = clean_cli_path(args.input).resolve()
            output_path = clean_cli_path(args.output).resolve()
            resource_paths = [clean_cli_path(path) for path in args.resource_path]
            result = convert_markdown(
                input_path,
                output_path,
                format_name=args.format,
                toc=args.toc,
                number_headings=args.number_headings,
                toc_depth=args.toc_depth,
                resource_paths=resource_paths,
            )
        else:  # pragma: no cover - argparse prevents this path
            parser.error(f"Unknown command: {args.command}")
            return 2
    except Exception as exc:
        if getattr(args, "json", False):
            print(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False), file=sys.stderr)
        else:
            print(f"Error: {exc}", file=sys.stderr)
        return 1

    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print(f"Wrote {result['output']}")
        warnings = result.get("warnings") or []
        for warning in warnings:
            print(f"Warning: {warning}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
